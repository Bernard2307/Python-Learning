from openpyxl import load_workbook
from docx import Document
from datetime import datetime

#INICIO

#carregar a planilha 
planilha = load_workbook('./W2 - CONTRATO AUTOMATICO/fornecedores.xlsx')
#passar a página da planilha 
planilha_fornecedores = planilha['Planilha1']

#PRINCIPAL

#criar o for para percorrer as linhas da planilha
for linha in planilha_fornecedores.iter_rows(min_row=2, values_only=True):
    nome_fornecedor, CNPJ, endereco, produto, preco, pagamento = linha #unpacking

    #criar o docx word para armazenar os contratos
    arqv_word = Document()
    arqv_word.add_heading("Contrato de prestação de serviços")
    
    #aplicar o texto do contrato
    txt_contrato = f"""
    **Contrato de Fornecimento**

    Este contrato é celebrado em {datetime.now().strftime('%d/%m/%Y')}, entre:

    **FORNECEDOR:**
    Nome: {nome_fornecedor}
    CNPJ: {CNPJ}
    Localidade: {endereco}

    E

    **COMPRADOR:**
    Nome: BLDAN LTDA AND COMPANY
    Endereço: Rua Joaquim Bispo, 178

    **Objeto do Contrato:**

    O Fornecedor concorda em fornecer os seguintes produtos ao Comprador:

    - Tipo de Produto: {produto}

    **Termos e Condições:**

    1. **Preço:** O preço total acordado para os produtos é de {preco}. O pagamento será efetuado de acordo com os termos estabelecidos pelas partes.

    4. **Condições de Pagamento:** As condições de pagamento serão {pagamento}.

    As partes concordam com os termos e condições acima mediante a assinatura deste contrato.
    """

    #adicionar o texto editado no word com um paragrafo 
    arqv_word.add_paragraph(txt_contrato)
    arqv_word.save(f'./W2 - CONTRATO AUTOMATICO/contratos/contrato{nome_fornecedor}.docx')

#salvar o arquivo em uma pasta escrito "contrato + nome da empresa"


